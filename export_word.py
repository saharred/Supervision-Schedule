"""
Word Document Export for Supervision Schedule
Matches the official template format
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#000000"},
        bottom={"sz": 12, "val": "single", "color": "#000000"},
        start={"sz": 12, "val": "single", "color": "#000000"},
        end={"sz": 12, "val": "single", "color": "#000000"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcPr.append(element)


def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)


def export_to_word(daily_schedule, school_name="مدرسة عثمان بن عفان النموذجية للبنين", 
                   academic_year="2025-2026", level_name="", semester="الفصل الدراسي الأول"):
    """
    Export daily schedule to Word document matching the template
    
    Args:
        daily_schedule: dict with date, day_name, and exams list
        school_name: name of the school
        academic_year: academic year
        level_name: level/grade name
        semester: semester name
    
    Returns:
        BytesIO buffer with Word document
    """
    doc = Document()
    
    # Set RTL for the document
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    
    # Header with logo and school info
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Right cell - School name in Arabic
    cell_right = header_table.rows[0].cells[2]
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(school_name)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run = p.add_run('\n')
    run = p.add_run('Othman bin Affan school for boys')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    
    # Middle cell - Academic year
    cell_middle = header_table.rows[0].cells[1]
    p = cell_middle.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'العام الأكاديمي {academic_year}')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Left cell - Ministry logo (placeholder text)
    cell_left = header_table.rows[0].cells[0]
    p = cell_left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('وزارة التربية والتعليم والتعليم العالي\n')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run = p.add_run('Ministry of Education and Higher Education')
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    
    doc.add_paragraph()
    
    # Title box
    title_table = doc.add_table(rows=3, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Row 1: Main title
    cell = title_table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('جدول مراقبة الورقة الأولى لاختبارات')
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Row 2: Semester
    cell = title_table.rows[1].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'منتصف {semester}')
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    
    # Row 3: Level/Period
    cell = title_table.rows[2].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'الورقة الأولى ({level_name})')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Add borders to title table
    for row in title_table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "#000000"},
                bottom={"sz": 12, "val": "single", "color": "#000000"},
                start={"sz": 12, "val": "single", "color": "#000000"},
                end={"sz": 12, "val": "single", "color": "#000000"}
            )
    
    doc.add_paragraph()
    
    # Date and subject info
    date_str = daily_schedule['date'].strftime('%Y-%m-%d')
    day_name = daily_schedule['day_name']
    
    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Row 1: Date
    cell = info_table.rows[0].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'اليوم: {day_name}')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    set_cell_background(cell, 'D9D9D9')
    
    cell = info_table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'التاريخ: {date_str}')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    set_cell_background(cell, 'D9D9D9')
    
    # Row 2: Subjects
    # Get unique subjects from exams
    subjects = set([exam['subject'] for exam in daily_schedule['exams']])
    subjects_str = ' - '.join(subjects)
    
    cell = info_table.rows[1].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'المادة: {subjects_str}')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    set_cell_background(cell, 'D9D9D9')
    
    # Get unique grades
    grades = set([exam['grade'] for exam in daily_schedule['exams']])
    grades_str = ' و '.join(grades)
    
    cell = info_table.rows[1].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'الصف: {grades_str}')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    set_cell_background(cell, 'D9D9D9')
    
    # Add borders
    for row in info_table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "#000000"},
                bottom={"sz": 12, "val": "single", "color": "#000000"},
                start={"sz": 12, "val": "single", "color": "#000000"},
                end={"sz": 12, "val": "single", "color": "#000000"}
            )
    
    doc.add_paragraph()
    
    # Main schedule table
    # Group exams by subject
    exams_by_subject = {}
    for exam in daily_schedule['exams']:
        subject = exam['subject']
        if subject not in exams_by_subject:
            exams_by_subject[subject] = []
        exams_by_subject[subject].append(exam)
    
    # Calculate total rows needed
    total_rows = sum(len(exams) for exams in exams_by_subject.values()) + len(exams_by_subject) + 2  # +2 for header and reserve
    
    # Create table: 6 columns (المادة، الصف، المراقب الأول، التوقيع، المراقب الثاني، التوقيع)
    schedule_table = doc.add_table(rows=total_rows, cols=6)
    schedule_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    headers = ['التوقيع', 'المراقب الثاني', 'التوقيع', 'المراقب الأول', 'الصف', 'المادة']
    for idx, header in enumerate(headers):
        cell = schedule_table.rows[0].cells[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '8B0000')  # Dark red
    
    # Fill data rows
    current_row = 1
    for subject, exams in exams_by_subject.items():
        # Merge cells for subject
        start_row = current_row
        
        for exam in exams:
            row = schedule_table.rows[current_row]
            
            # Column 5: الصف
            cell = row.cells[4]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{exam['grade']} {exam['section']}")
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            
            # Column 4: المراقب الأول
            cell = row.cells[3]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(exam['supervisor1'])
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            
            # Column 3: التوقيع (empty)
            cell = row.cells[2]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Column 2: المراقب الثاني
            cell = row.cells[1]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(exam['supervisor2'])
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            
            # Column 1: التوقيع (empty)
            cell = row.cells[0]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            current_row += 1
        
        # Merge subject cells
        end_row = current_row - 1
        if start_row <= end_row:
            merged_cell = schedule_table.rows[start_row].cells[5].merge(schedule_table.rows[end_row].cells[5])
            p = merged_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(subject)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            set_cell_background(merged_cell, 'FFC7CE')  # Light pink
    
    # Reserve row (الاحتياط)
    row = schedule_table.rows[current_row]
    merged_cell = row.cells[0].merge(row.cells[5])
    p = merged_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('الاحتياط')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_background(merged_cell, '8B0000')
    
    # Add borders to all cells
    for row in schedule_table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "#000000"},
                bottom={"sz": 12, "val": "single", "color": "#000000"},
                start={"sz": 12, "val": "single", "color": "#000000"},
                end={"sz": 12, "val": "single", "color": "#000000"}
            )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer with signatures
    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    cell = footer_table.rows[0].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('النائب الأكاديمي: مريم القصع')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    
    cell = footer_table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('النائب الإداري: دلال الفهيدة')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # School principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('مديرة المدرسة: منيرة الهاجري')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Motto
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('رؤيتنا: مجتمع رياضي لتنمية مستدامة')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.italic = True
    
    # Save to file
    import tempfile
    import os
    
    # Create temp file
    fd, temp_path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    
    # Save document
    doc.save(temp_path)
    
    return temp_path

